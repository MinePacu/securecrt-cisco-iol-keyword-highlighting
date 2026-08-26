from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "CRT_Cisco_IOL_Highlight_설치_안내서.docx"

# compact_reference_guide preset tokens
PAGE_WIDTH_DXA = 12240  # Letter, 8.5 in
PAGE_HEIGHT_DXA = 15840  # Letter, 11 in
MARGIN_DXA = 1440  # 1 in
CONTENT_WIDTH_DXA = 9360  # 6.5 in
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
LINE_125 = 1.25

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F7FB"
LIGHT_GRAY = "F2F4F7"
FRAME_BORDER = "9BAEC2"
CAUTION_FILL = "FFF7E6"
CAUTION_BORDER = "D6A642"
SUCCESS_FILL = "EEF7F0"
SUCCESS_BORDER = "7BA982"
BLACK = "1F2933"
# Apple SD Gothic Neo is installed in the render environment
# (AppleSDGothicNeo.ttc) and renders Korean glyphs reliably in LibreOffice
# headless conversion.
RENDER_EAST_ASIA = "Apple SD Gothic Neo"


def set_xml_attr(element, name: str, value: str) -> None:
    element.set(qn(name), value)


def ensure_rfonts(parent) -> object:
    rpr = parent.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        parent.insert(0, rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    return rfonts


def set_style_font(style, latin: str = "Calibri", east_asia: str = RENDER_EAST_ASIA) -> None:
    style.font.name = latin
    rfonts = ensure_rfonts(style._element)
    set_xml_attr(rfonts, "w:ascii", latin)
    set_xml_attr(rfonts, "w:hAnsi", latin)
    set_xml_attr(rfonts, "w:eastAsia", east_asia)
    set_xml_attr(rfonts, "w:cs", latin)
    set_xml_attr(rfonts, "w:hint", "eastAsia")


def set_run_font(
    run,
    latin: str = "Calibri",
    east_asia: str = RENDER_EAST_ASIA,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    # Use the verified installed CJK font for every direct run so the
    # headless renderer cannot substitute a font without Korean glyphs.
    run.font.name = RENDER_EAST_ASIA
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    set_xml_attr(rfonts, "w:ascii", RENDER_EAST_ASIA)
    set_xml_attr(rfonts, "w:hAnsi", RENDER_EAST_ASIA)
    set_xml_attr(rfonts, "w:eastAsia", RENDER_EAST_ASIA)
    set_xml_attr(rfonts, "w:cs", RENDER_EAST_ASIA)
    set_xml_attr(rfonts, "w:hint", "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    set_xml_attr(lang, "w:val", "ko-KR")
    set_xml_attr(lang, "w:eastAsia", "ko-KR")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    set_xml_attr(shd, "w:val", "clear")
    set_xml_attr(shd, "w:color", "auto")
    set_xml_attr(shd, "w:fill", fill)


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in margins.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        set_xml_attr(node, "w:w", str(value))
        set_xml_attr(node, "w:type", "dxa")


def set_cell_borders(cell, color: str = FRAME_BORDER, size: int = 8, val: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        set_xml_attr(node, "w:val", val)
        set_xml_attr(node, "w:sz", str(size))
        set_xml_attr(node, "w:space", "0")
        set_xml_attr(node, "w:color", color)


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    """Apply fixed DXA table geometry required by the document preset."""
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA} DXA: {widths}")
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    set_xml_attr(tbl_w, "w:type", "dxa")
    set_xml_attr(tbl_w, "w:w", str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    set_xml_attr(tbl_ind, "w:type", "dxa")
    set_xml_attr(tbl_ind, "w:w", str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    set_xml_attr(layout, "w:type", "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        set_xml_attr(col, "w:w", str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            set_xml_attr(tc_w, "w:type", "dxa")
            set_xml_attr(tc_w, "w:w", str(widths[idx]))
            set_cell_margins(cell)


def set_min_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    set_xml_attr(tr_height, "w:val", str(height_twips))
    set_xml_attr(tr_height, "w:hRule", "atLeast")
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_paragraph_border(paragraph, color: str = BLUE, size: int = 12, space: int = 5) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    set_xml_attr(bottom, "w:val", "single")
    set_xml_attr(bottom, "w:sz", str(size))
    set_xml_attr(bottom, "w:space", str(space))
    set_xml_attr(bottom, "w:color", color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    set_xml_attr(fld_begin, "w:fldCharType", "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    set_xml_attr(instr, "xml:space", "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    set_xml_attr(fld_sep, "w:fldCharType", "separate")
    run._r.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    set_xml_attr(fld_end, "w:fldCharType", "end")
    run._r.append(fld_end)


def configure_numbering(document: Document) -> dict[str, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract + [0]) + 1
    num_id = max(existing_nums + [0]) + 1

    def add_abstract(kind: str, marker: str) -> int:
        nonlocal abstract_id
        aid = abstract_id
        abstract_id += 1
        abstract = OxmlElement("w:abstractNum")
        set_xml_attr(abstract, "w:abstractNumId", str(aid))
        multi = OxmlElement("w:multiLevelType")
        set_xml_attr(multi, "w:val", "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        set_xml_attr(lvl, "w:ilvl", "0")
        start = OxmlElement("w:start")
        set_xml_attr(start, "w:val", "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        set_xml_attr(fmt, "w:val", kind)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        set_xml_attr(lvl_text, "w:val", marker)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        set_xml_attr(suff, "w:val", "tab")
        lvl.append(suff)
        lvl_jc = OxmlElement("w:lvlJc")
        set_xml_attr(lvl_jc, "w:val", "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        set_xml_attr(tab, "w:val", "num")
        set_xml_attr(tab, "w:pos", "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        set_xml_attr(ind, "w:left", "540")
        set_xml_attr(ind, "w:hanging", "270")
        p_pr.append(ind)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        set_xml_attr(r_fonts, "w:ascii", "Calibri")
        set_xml_attr(r_fonts, "w:hAnsi", "Calibri")
        set_xml_attr(r_fonts, "w:eastAsia", "Malgun Gothic")
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        return aid

    def add_num(aid: int) -> int:
        nonlocal num_id
        nid = num_id
        num_id += 1
        num = OxmlElement("w:num")
        set_xml_attr(num, "w:numId", str(nid))
        abs_id = OxmlElement("w:abstractNumId")
        set_xml_attr(abs_id, "w:val", str(aid))
        num.append(abs_id)
        numbering.append(num)
        return nid

    bullet_abstract = add_abstract("bullet", "•")
    decimal_abstract = add_abstract("decimal", "%1.")
    check_abstract = add_abstract("bullet", "☐")
    return {
        "bullet": add_num(bullet_abstract),
        "bullet_abstract": bullet_abstract,
        "decimal": add_num(decimal_abstract),
        "decimal_abstract": decimal_abstract,
        "check": add_num(check_abstract),
        "check_abstract": check_abstract,
    }


def new_num_instance(document: Document, abstract_id: int) -> int:
    """Create a real Word numbering instance that restarts at 1."""
    numbering = document.part.numbering_part.element
    existing_nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums + [0]) + 1
    num = OxmlElement("w:num")
    set_xml_attr(num, "w:numId", str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    set_xml_attr(abs_id, "w:val", str(abstract_id))
    num.append(abs_id)
    override = OxmlElement("w:lvlOverride")
    set_xml_attr(override, "w:ilvl", "0")
    start = OxmlElement("w:startOverride")
    set_xml_attr(start, "w:val", "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    set_xml_attr(ilvl, "w:val", str(level))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    set_xml_attr(num_id_node, "w:val", str(num_id))


def set_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = LINE_125

    def style(name: str, base: str = "Normal"):
        if name in styles:
            s = styles[name]
        else:
            s = styles.add_style(name, 1)
        s.base_style = styles[base]
        return s

    title = style("Guide Title")
    set_style_font(title)
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK_BLUE)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.05
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = style("Guide Subtitle")
    set_style_font(subtitle)
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.line_spacing = 1.15
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    kicker = style("Cover Kicker")
    set_style_font(kicker)
    kicker.font.size = Pt(10.5)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(BLUE)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(16)
    kicker.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        s = styles[name]
        set_style_font(s)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.line_spacing = LINE_125
        s.paragraph_format.keep_with_next = True

    list_style = style("Guide List")
    set_style_font(list_style)
    list_style.font.size = Pt(11)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = LINE_125

    small = style("Small Note")
    set_style_font(small)
    small.font.size = Pt(9.5)
    small.font.color.rgb = RGBColor.from_string(MUTED)
    small.paragraph_format.space_before = Pt(0)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.2

    caption = style("Figure Caption")
    set_style_font(caption)
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    caption.font.bold = True
    caption.paragraph_format.space_before = Pt(5)
    caption.paragraph_format.space_after = Pt(1)
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    figure_action = style("Figure Action")
    set_style_font(figure_action)
    figure_action.font.size = Pt(9.5)
    figure_action.font.color.rgb = RGBColor.from_string(MUTED)
    figure_action.paragraph_format.space_before = Pt(0)
    figure_action.paragraph_format.space_after = Pt(9)
    figure_action.paragraph_format.line_spacing = 1.15
    figure_action.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callout_title = style("Callout Title")
    set_style_font(callout_title)
    callout_title.font.size = Pt(11)
    callout_title.font.bold = True
    callout_title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout_title.paragraph_format.space_before = Pt(0)
    callout_title.paragraph_format.space_after = Pt(3)
    callout_title.paragraph_format.line_spacing = 1.15

    callout_body = style("Callout Body")
    set_style_font(callout_body)
    callout_body.font.size = Pt(10.5)
    callout_body.paragraph_format.space_before = Pt(0)
    callout_body.paragraph_format.space_after = Pt(0)
    callout_body.paragraph_format.line_spacing = 1.2


def add_text_runs(paragraph, segments: Iterable[tuple[str, dict]]) -> None:
    for text, attrs in segments:
        run = paragraph.add_run(text)
        if attrs.get("code"):
            set_run_font(run, latin="Consolas", east_asia="Malgun Gothic", size=10.5, color=INK_BLUE)
            run.bold = attrs.get("bold", False)
        else:
            set_run_font(
                run,
                size=attrs.get("size"),
                color=attrs.get("color"),
                bold=attrs.get("bold"),
                italic=attrs.get("italic"),
            )


def add_para(document, segments: Sequence[tuple[str, dict]], style: str = "Normal", align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    add_text_runs(paragraph, segments)
    return paragraph


def add_list_item(document, num_id: int, segments: Sequence[tuple[str, dict]], style: str = "Guide List"):
    paragraph = document.add_paragraph(style=style)
    apply_numbering(paragraph, num_id)
    add_text_runs(paragraph, segments)
    return paragraph


def add_spacer(document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = 1


def add_heading(document, text: str, level: int = 1):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    for run in paragraph.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return paragraph


def add_callout(
    document,
    title: str,
    body: Sequence[tuple[str, dict]],
    fill: str = LIGHT_BLUE,
    border: str = FRAME_BORDER,
) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, border, size=10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].style = "Callout Title"
    cell.paragraphs[0].add_run(title)
    set_run_font(cell.paragraphs[0].runs[0], size=11, color=DARK_BLUE, bold=True)
    paragraph = cell.add_paragraph(style="Callout Body")
    add_text_runs(paragraph, body)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)
    spacer.paragraph_format.line_spacing = 1


def add_placeholder(
    document,
    number: int,
    target: str,
    caption: str,
    action: str,
    height_twips: int = 2500,
) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EDF3F8")
    set_cell_borders(cell, FRAME_BORDER, size=14)
    set_min_row_height(table.rows[0], height_twips)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"사진 {number} 삽입 위치")
    set_run_font(r, size=18, color=FRAME_BORDER, bold=True)
    p2 = cell.add_paragraph(style="Small Note")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run(f"촬영 대상: {target}")
    set_run_font(r, size=9.5, color=MUTED, bold=True)
    p3 = cell.add_paragraph(style="Small Note")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r = p3.add_run("이 안내 문구를 모두 지운 뒤 사진을 삽입하세요.")
    set_run_font(r, size=9.5, color=MUTED)
    cap = document.add_paragraph(style="Figure Caption")
    cap.add_run(caption)
    set_run_font(cap.runs[0], size=9.5, color=DARK_BLUE, bold=True)
    act = document.add_paragraph(style="Figure Action")
    act.add_run(action)
    set_run_font(act.runs[0], size=9.5, color=MUTED)


def add_code_run_paragraph(document, prefix: str, code: str, suffix: str = "", style: str = "Normal"):
    return add_para(
        document,
        [(prefix, {}), (code, {"code": True, "bold": True}), (suffix, {})],
        style=style,
    )


def configure_header_footer(document: Document) -> None:
    for section in document.sections:
        set_section_geometry(section)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(2)
        hp.paragraph_format.line_spacing = 1
        run = hp.add_run("CRT Cisco IOL Highlight 설치 안내서")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        set_paragraph_border(hp, color="C7D2DE", size=6, space=3)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(2)
        fp.paragraph_format.line_spacing = 1
        label = fp.add_run("페이지 ")
        set_run_font(label, size=8.5, color=MUTED)
        add_page_field(fp)


def build_document() -> Document:
    document = Document()
    document.core_properties.title = "CRT Cisco IOL Highlight 설치 안내서"
    document.core_properties.subject = "Windows에서 SecureCRT Cisco IOL 키워드 하이라이트 설치하기"
    document.core_properties.author = "CRT Cisco IOL Highlight"
    document.core_properties.keywords = "SecureCRT, Cisco IOL, Windows, 설치 안내서"

    configure_styles(document)
    numbering = configure_numbering(document)
    configure_header_footer(document)

    # Page 1 — editorial_cover style: simple, centered, generous title block.
    add_spacer(document, 48)
    add_para(document, [("WINDOWS 설치 가이드", {"bold": True, "color": BLUE, "size": 10.5})], style="Cover Kicker", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(document, [("CRT Cisco IOL Highlight 설치 안내서", {"bold": True, "color": INK_BLUE, "size": 29})], style="Guide Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(document, [("SecureCRT에 Cisco IOL 키워드 하이라이트 적용하기", {"color": BLUE, "size": 15})], style="Guide Subtitle", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_spacer(document, 10)
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(16)
    set_paragraph_border(rule, color=BLUE, size=12, space=2)
    add_para(
        document,
        [("이 안내서는 Windows에서 배포 ZIP을 풀고 설치 파일을 실행해 SecureCRT에 하이라이트를 적용하는 과정을 설명합니다. 사진만 나중에 추가하면 배포용 안내서로 사용할 수 있습니다.", {})],
        style="Normal",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_heading(document, "사진을 추가하는 방법", 2)
    add_list_item(document, numbering["bullet"], [("각 회색·연한 파랑 프레임은 실제 화면 사진을 넣을 자리입니다.", {})])
    add_list_item(document, numbering["bullet"], [("사진을 넣을 때는 프레임 안의 안내 문구를 모두 지운 뒤 사진을 삽입하세요.", {})])
    add_list_item(document, numbering["bullet"], [("특히 SmartScreen의 두 화면은 순서가 중요합니다. 먼저 ", {}), ("추가 정보", {"bold": True}), ("를 누르는 화면과, 확장된 뒤 ", {}), ("계속 실행", {"bold": True}), (" 또는 ", {}), ("실행", {"bold": True}), ("을 누르는 화면을 각각 넣어 주세요.", {})])
    add_callout(
        document,
        "이 문서에서 꼭 기억할 것",
        [("설치 파일만 따로 옮기지 말고, 배포 ZIP을 먼저 전체 압축 해제한 폴더에서 실행합니다. 설치 중 확인 질문에는 ", {}), ("Y", {"code": True, "bold": True}), (" 또는 ", {}), ("y", {"code": True, "bold": True}), ("만 승인으로 인식됩니다.", {})],
        fill=PALE_BLUE,
        border=FRAME_BORDER,
    )
    add_para(document, [("대상 환경: Windows · SecureCRT · Cisco IOL CLI 출력", {"color": MUTED, "size": 9.5, "italic": True})], style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)

    document.add_page_break()

    # Page 2 — preparation and quick flow.
    add_heading(document, "설치 전 체크리스트", 1)
    add_para(document, [("설치를 시작하기 전에 아래 항목을 확인하면 중간에 멈추는 일을 줄일 수 있습니다.", {})])
    page2_check_num = new_num_instance(document, numbering["check_abstract"])
    for item in [
        "Windows 컴퓨터에서 작업하고 있는지 확인합니다.",
        "설치 전 SecureCRT를 완전히 종료합니다. 실행 중이면 설치 스크립트가 계속할지 물을 수 있습니다.",
        "배포 ZIP 전체를 받을 수 있는지 확인합니다. ZIP 안의 파일을 함께 보존해야 합니다.",
        "기존 설정을 별도로 보관하고 싶다면 SecureCRT 설정 폴더를 추가로 백업합니다. 설치 프로그램도 덮어쓰기 직전에 timestamp 백업을 만듭니다.",
    ]:
        add_list_item(document, page2_check_num, [(item, {})])
    add_heading(document, "설치 흐름 한눈에 보기", 1)
    steps = [
        "배포 ZIP을 다운로드합니다.",
        "ZIP에서 ‘압축 풀기/Extract All’을 선택해 전체 압축을 풉니다.",
        "압축 해제 폴더에서 Install-KeywordHighlight-Setup.exe를 더블클릭합니다.",
        "SmartScreen이 보이면 추가 정보 → 계속 실행/실행 순서로 진행합니다.",
        "PowerShell 설치 창의 질문에는 Y 후 Enter를 입력합니다.",
        "설치가 끝나면 SecureCRT를 다시 시작하고 테스트 출력을 확인합니다.",
    ]
    page2_decimal_num = new_num_instance(document, numbering["decimal_abstract"])
    for step in steps:
        add_list_item(document, page2_decimal_num, [(step, {})])
    add_callout(
        document,
        "이 안내서에 나오는 문자열",
        [("파일 이름: ", {"bold": True}), ("Install-KeywordHighlight-Setup.exe", {"code": True, "bold": True}), (" · 키워드 파일: ", {"bold": True}), ("Keywords\\PNET-Cisco-Dark.ini", {"code": True, "bold": True}), (" · 기본 세션 파일: ", {"bold": True}), ("Sessions\\Default.ini", {"code": True, "bold": True}), (" · 승인 입력: ", {"bold": True}), ("Y", {"code": True, "bold": True}), (" 또는 ", {}), ("y", {"code": True, "bold": True})],
        fill=LIGHT_GRAY,
        border="B8C4D0",
    )
    add_para(document, [("PowerShell 명령을 직접 입력할 필요 없이, 아래에서는 배포된 원클릭 실행 파일 흐름을 기준으로 설명합니다.", {"italic": True, "color": MUTED})], style="Small Note")

    document.add_page_break()

    # Page 3 — distribution package and executable.
    add_heading(document, "1. 배포 ZIP을 다운로드하고 실행 파일 열기", 1)
    add_para(document, [("배포본은 실행 파일과 함께 필요한 파일이 들어 있는 폴더 또는 ZIP 전체를 배포하는 방식입니다. ZIP을 풀지 않고 실행하거나 실행 파일만 옮기면 설치에 필요한 파일을 찾지 못할 수 있습니다.", {})])
    page3_decimal_num = new_num_instance(document, numbering["decimal_abstract"])
    add_list_item(document, page3_decimal_num, [("배포 ZIP을 다운로드합니다. 다운로드가 끝날 때까지 기다립니다.", {})])
    add_list_item(document, page3_decimal_num, [("다운로드한 ZIP을 마우스 오른쪽 버튼으로 클릭하고 ", {}), ("압축 풀기/Extract All", {"bold": True}), ("을 선택합니다. 파일 전체가 새 폴더에 풀리도록 합니다.", {})])
    add_list_item(document, page3_decimal_num, [("압축 해제된 폴더를 열고 ", {}), ("Install-KeywordHighlight-Setup.exe", {"code": True, "bold": True}), ("를 더블클릭합니다.", {})])
    add_callout(
        document,
        "파일을 옮기지 마세요",
        [("설치가 끝날 때까지 실행 파일만 다른 폴더로 옮기지 마세요. ZIP에서 함께 풀린 파일과 같은 폴더에 둔 상태로 실행해야 합니다.", {})],
        fill=CAUTION_FILL,
        border=CAUTION_BORDER,
    )
    add_placeholder(
        document,
        1,
        "다운로드한 ZIP 파일과 마우스 오른쪽 버튼 메뉴",
        "사진 1 — ZIP에서 ‘압축 풀기/Extract All’을 선택하는 화면",
        "행동: ZIP 전체를 압축 해제한 뒤 새 폴더를 엽니다.",
        height_twips=1750,
    )
    add_placeholder(
        document,
        2,
        "압축 해제 폴더 안의 Install-KeywordHighlight-Setup.exe",
        "사진 2 — 압축 해제 폴더에서 설치 실행 파일을 찾은 화면",
        "행동: 파일을 폴더 밖으로 옮기지 말고 더블클릭합니다.",
        height_twips=1750,
    )

    document.add_page_break()

    # Page 4 — SmartScreen, intentionally two large placeholders.
    add_heading(document, "2. Windows Defender SmartScreen에서 계속 실행", 1)
    add_callout(
        document,
        "중요: ‘계속 실행’ 버튼이 나오는 과정",
        [("SmartScreen 화면이 보이면 전체 SmartScreen을 해제하지 마세요. 배포 출처와 파일 이름을 확인한 뒤, 먼저 ", {}), ("추가 정보", {"bold": True}), ("를 누르고, 확장된 화면에서 ", {}), ("계속 실행", {"bold": True}), (" 또는 Windows 버전에 따라 ", {}), ("실행", {"bold": True}), (" 버튼을 누릅니다.", {})],
        fill=CAUTION_FILL,
        border=CAUTION_BORDER,
    )
    page4_decimal_num = new_num_instance(document, numbering["decimal_abstract"])
    add_list_item(document, page4_decimal_num, [("처음 표시된 SmartScreen 창에서 배포 출처와 파일 이름이 예상한 것과 맞는지 확인합니다.", {})])
    add_list_item(document, page4_decimal_num, [("화면에 ", {}), ("추가 정보", {"bold": True}), ("가 보이면 먼저 클릭합니다.", {})])
    add_list_item(document, page4_decimal_num, [("내용이 펼쳐지면 ", {}), ("계속 실행", {"bold": True}), ("을 클릭합니다. Windows 버전에 따라 버튼이 ", {}), ("실행", {"bold": True}), ("으로 표시될 수 있습니다.", {})])
    add_placeholder(
        document,
        3,
        "SmartScreen 처음 표시된 화면의 추가 정보 버튼",
        "사진 3 — SmartScreen 처음 화면: ‘추가 정보’가 보이는 상태",
        "행동: 먼저 ‘추가 정보’를 눌러 다음 선택지를 표시합니다.",
        height_twips=2400,
    )
    add_placeholder(
        document,
        4,
        "SmartScreen 확장 화면의 계속 실행/실행 버튼",
        "사진 4 — SmartScreen 확장 화면: ‘계속 실행’ 또는 ‘실행’ 버튼",
        "행동: 출처를 확인한 뒤 ‘계속 실행’ 또는 ‘실행’을 눌러 설치를 시작합니다.",
        height_twips=2400,
    )

    document.add_page_break()

    # Page 5 — UAC and PowerShell prompts.
    add_heading(document, "3. UAC와 PowerShell 설치 창에서 진행하기", 1)
    add_heading(document, "UAC 권한 확인", 2)
    add_para(document, [("Windows 사용자 계정 컨트롤(UAC) 창이 표시될 때만 진행합니다. 창에 표시된 게시자·출처가 배포한 파일과 일치하는지 확인한 뒤 ", {}), ("예", {"bold": True}), ("를 누릅니다. 출처가 예상과 다르면 설치를 중단하고 배포 담당자에게 확인하세요.", {})])
    add_heading(document, "PowerShell 창의 질문에 답하기", 2)
    page5_decimal_num = new_num_instance(document, numbering["decimal_abstract"])
    for item in [
        [("PowerShell 설치 창이 열리면 창을 닫지 말고 작업이 진행될 때까지 기다립니다.", {})],
        [("SecureCRT가 실행 중이라는 확인 또는 기존 ", {}), ("PNET-Cisco-Dark.ini", {"code": True}), ("/", {}), ("Default.ini", {"code": True}), (" 적용 확인이 나오면 ", {}), ("Y", {"code": True, "bold": True}), (" 또는 ", {}), ("y", {"code": True, "bold": True}), ("를 입력하고 Enter를 누릅니다. ", {}), ("yes", {"code": True}), ("나 ", {}), ("예", {"code": True}), ("는 승인 입력이 아닙니다.", {})],
        [("설정 폴더 자동 탐색에 실패하면, ", {}), ("Sessions\\Default.ini", {"code": True, "bold": True}), ("가 들어 있는 SecureCRT 설정 폴더의 경로를 입력하고 Enter를 누릅니다. 파일 자체가 아니라 그 상위 설정 폴더 경로를 입력합니다.", {})],
        [("완료 메시지와 ", {}), ("Press any key to continue...", {"code": True, "bold": True}), ("가 보이면 아무 키나 눌러 창을 닫습니다.", {})],
    ]:
        add_list_item(document, page5_decimal_num, item)
    add_callout(
        document,
        "설치 파일이 하는 일",
        [("설치 프로그램은 ", {}), ("Keywords\\PNET-Cisco-Dark.ini", {"code": True, "bold": True}), ("를 적용하고 ", {}), ("Sessions\\Default.ini", {"code": True, "bold": True}), ("의 기본 세션 옵션을 적용합니다. 기존 파일을 덮어쓰기 전에는 timestamp 백업을 만듭니다.", {})],
        fill=SUCCESS_FILL,
        border=SUCCESS_BORDER,
    )
    add_placeholder(
        document,
        5,
        "PowerShell 설치 창, Y/N 질문 또는 완료 메시지",
        "사진 5 — PowerShell 창에서 Y 입력 또는 완료 메시지를 확인하는 화면",
        "행동: 질문에는 Y 후 Enter를 입력하고, 완료 후 아무 키나 눌러 창을 닫습니다.",
        height_twips=2400,
    )

    document.add_page_break()

    # Page 6 — restart and troubleshooting.
    add_heading(document, "4. SecureCRT를 다시 시작하고 색상 확인하기", 1)
    page6_decimal_num = new_num_instance(document, numbering["decimal_abstract"])
    add_list_item(document, page6_decimal_num, [("설치 창을 닫은 뒤 SecureCRT를 다시 시작합니다.", {})])
    add_list_item(document, page6_decimal_num, [("Cisco IOL 장비에 연결하거나 저장된 CLI 출력으로 아래 명령 결과를 확인합니다.", {})])
    add_list_item(document, page6_decimal_num, [("show interfaces", {"code": True, "bold": True}), (" · ", {}), ("show spanning-tree", {"code": True, "bold": True}), (" · ", {}), ("show etherchannel summary", {"code": True, "bold": True}), (" · ", {}), ("show standby", {"code": True, "bold": True}), (" · ", {}), ("show ip ospf neighbor", {"code": True, "bold": True}), (" 중 하나를 실행하거나 출력합니다.", {})])
    add_list_item(document, page6_decimal_num, [("상태, 인터페이스, STP, EtherChannel, HSRP, OSPF 관련 문자열에 색상 하이라이트가 보이는지 확인합니다.", {})])
    add_heading(document, "문제 해결", 1)
    add_heading(document, "SmartScreen 버튼이 안 보일 때", 2)
    add_para(document, [("Windows 버전에 따라 화면 구성이 다를 수 있습니다. 창에 ", {}), ("추가 정보", {"bold": True}), ("가 보이는지 먼저 확인하고, 그 다음 ", {}), ("계속 실행", {"bold": True}), (" 또는 ", {}), ("실행", {"bold": True}), ("을 찾습니다. SmartScreen 전체 해제는 안내하지 않습니다. 파일을 받은 배포 출처와 파일 이름을 다시 확인하고, 의심스러우면 실행하지 마세요.", {})])
    add_heading(document, "설정 폴더를 찾지 못할 때", 2)
    add_para(document, [("자동 탐색이 실패하면 ", {}), ("Sessions\\Default.ini", {"code": True, "bold": True}), ("가 실제로 들어 있는 SecureCRT 설정 폴더 경로를 입력합니다. ", {}), ("Sessions", {"code": True}), (" 폴더나 ", {}), ("Default.ini", {"code": True}), (" 파일만 단독으로 입력하지 않습니다. 경로를 모르면 SecureCRT 설정 폴더를 먼저 확인하세요.", {})])
    add_heading(document, "색상이 안 보일 때", 2)
    add_para(document, [("SecureCRT가 설치 전에 열려 있었다면 완전히 종료한 뒤 다시 시작합니다. 다시 시작해도 보이지 않으면 하이라이트가 적용된 세션으로 연결했는지, 출력이 실제 Cisco IOL 형식인지, 터미널 테마와 색상이 충돌하지 않는지 확인합니다.", {})])
    add_heading(document, "압축을 풀지 않고 실행했을 때", 2)
    add_para(document, [("ZIP을 먼저 ", {}), ("압축 풀기/Extract All", {"bold": True}), ("로 전체 해제한 뒤, 압축 해제 폴더 안에서 ", {}), ("Install-KeywordHighlight-Setup.exe", {"code": True, "bold": True}), ("를 다시 실행합니다. 실행 파일만 다른 폴더로 옮기지 않습니다.", {})])

    # Page 7 — final acceptance checklist and last photo.
    add_heading(document, "설치가 끝났는지 확인", 1)
    add_para(document, [("아래 항목을 모두 확인하면 설치가 완료된 것입니다.", {})])
    page7_check_num = new_num_instance(document, numbering["check_abstract"])
    for item in [
        "배포 ZIP을 전체 압축 해제한 폴더에서 실행했습니다.",
        "SmartScreen에서 추가 정보 → 계속 실행/실행 순서로 진행했습니다.",
        "UAC가 표시되었을 때 출처를 확인하고 예를 눌렀습니다.",
        "설치 확인 질문과 기존 파일 적용 확인에 Y 또는 y를 입력했습니다.",
        "설정 폴더 자동 탐색 실패 시 Sessions\\Default.ini가 들어 있는 설정 폴더 경로를 입력했습니다.",
        "설치 완료 메시지와 Press any key to continue가 보인 뒤 키를 눌러 창을 닫았습니다.",
        "SecureCRT를 다시 시작했고 Cisco IOL 출력에서 색상 하이라이트를 확인했습니다.",
    ]:
        add_list_item(document, page7_check_num, [(item, {})])
    add_callout(
        document,
        "수동 INI 복사는 필요하지 않습니다",
        [("원클릭 설치 프로그램이 ", {}), ("Keywords\\PNET-Cisco-Dark.ini", {"code": True, "bold": True}), ("와 ", {}), ("Sessions\\Default.ini", {"code": True, "bold": True}), ("의 기본 세션 옵션을 함께 적용합니다. 설치 후 INI 파일을 별도로 복사하거나 세션 옵션을 수동으로 입력하지 않아도 됩니다.", {})],
        fill=SUCCESS_FILL,
        border=SUCCESS_BORDER,
    )
    add_placeholder(
        document,
        6,
        "SecureCRT 재시작 후 Cisco IOL 출력에 색상이 표시된 화면",
        "사진 6 — 설치 후 SecureCRT에서 Cisco IOL 출력 하이라이트를 확인하는 화면",
        "행동: 실제 배포본에서 색상 하이라이트가 보이는 최종 결과 사진을 넣습니다.",
        height_twips=2950,
    )
    add_para(document, [("끝. 이 문서는 사진 1~6의 회색 안내 문구를 지운 뒤 실제 스크린샷으로 교체해 배포하세요.", {"italic": True, "color": MUTED, "size": 9.5})], style="Small Note", align=WD_ALIGN_PARAGRAPH.CENTER)

    return document


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)
